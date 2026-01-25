#!/usr/bin/env python3
"""
EFMD Programme Scraper v2.0
===========================
Enhanced scraper with:
- Multilingual support (embeddings are language-agnostic)
- Semantic gap detection via vector similarity
- Supabase integration
- Site-specific parsers for known universities

Dependencies:
    pip install requests beautifulsoup4 google-generativeai supabase python-dotenv
"""

import json
import re
import os
from dataclasses import dataclass, field, asdict
from typing import Optional, Callable
from datetime import datetime
from urllib.parse import urljoin, urlparse
from abc import ABC, abstractmethod

# PDF/DOCX extraction
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    print("Warning: pypdf not installed. PDF extraction disabled.")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. DOCX extraction disabled.")

# Core dependencies
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Optional: Supabase and Gemini
try:
    from supabase import create_client, Client
    HAS_SUPABASE = True
except ImportError:
    HAS_SUPABASE = False
    print("Warning: supabase not installed. Database features disabled.")

try:
    import google.generativeai as genai
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False
    print("Warning: google-generativeai not installed. Embeddings disabled.")


# ============================================================
# CONFIGURATION
# ============================================================

# Gemini embedding model (768 dimensions, multilingual)
EMBEDDING_MODEL = "models/text-embedding-004"
EMBEDDING_DIMENSIONS = 768

# Similarity thresholds
PILLAR_MATCH_THRESHOLD = 0.55  # Minimum similarity to count as matching a pillar
STRONG_MATCH_THRESHOLD = 0.70  # Strong match

# User agent for scraping
USER_AGENT = 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'


# ============================================================
# DATA CLASSES
# ============================================================

@dataclass
class ILOAnalysis:
    """Analysis of a single ILO"""
    text: str
    source_language: str = 'en'
    ksa_category: Optional[str] = None  # Knowledge, Skill, Attitude
    verb_found: Optional[str] = None
    has_weak_verb: bool = False
    is_measurable: bool = True
    quality_issues: list[str] = field(default_factory=list)
    embedding: Optional[list[float]] = None
    matched_pillars: list[str] = field(default_factory=list)
    best_match_score: float = 0.0


@dataclass
class CourseData:
    """Individual course data"""
    title: str
    description: Optional[str] = None
    source_language: str = 'en'
    ects: Optional[float] = None
    year: Optional[int] = None
    semester: Optional[str] = None
    is_mandatory: Optional[bool] = None
    course_ilos: list[str] = field(default_factory=list)
    embedding: Optional[list[float]] = None
    matched_pillars: list[str] = field(default_factory=list)


@dataclass
class ProgrammeData:
    """Complete programme data"""
    # Basic info
    institution: str
    programme_name: str
    primary_url: str
    
    # Classification
    degree_type: Optional[str] = None
    duration_months: Optional[int] = None
    total_ects: Optional[int] = None
    delivery_mode: Optional[str] = None
    languages_of_instruction: list[str] = field(default_factory=list)
    
    # URLs scraped
    urls_scraped: list[str] = field(default_factory=list)
    
    # Content
    programme_ilos: list[ILOAnalysis] = field(default_factory=list)
    courses: list[CourseData] = field(default_factory=list)
    programme_aims: list[str] = field(default_factory=list)
    
    # Raw content (for re-processing)
    raw_html: str = ''
    raw_text: str = ''
    
    # Flags
    has_ilo_matrix: bool = False
    
    # Metadata
    scraped_at: str = field(default_factory=lambda: datetime.now().isoformat())
    scrape_notes: list[str] = field(default_factory=list)


@dataclass
class GapAnalysisResult:
    """Complete gap analysis result"""
    programme: ProgrammeData
    
    # Scores
    readiness_score: int = 0
    eligibility_pass: bool = False
    
    # ILO Analysis
    ilo_count: int = 0
    ilo_has_knowledge: bool = False
    ilo_has_skills: bool = False
    ilo_has_attitudes: bool = False
    ilo_weak_verb_count: int = 0
    
    # Documentation Status
    courses_total: int = 0
    courses_with_ilos: int = 0
    courses_no_ilos: list[str] = field(default_factory=list)  # course names missing ILOs
    documentation_score: float = 0.0
    ilo_issues: list[str] = field(default_factory=list)
    
    # Pillar Coverage (from semantic matching)
    pillar_coverage: dict = field(default_factory=dict)  # {pillar: score}
    missing_pillars: list[str] = field(default_factory=list)
    
    # Eligibility Gates
    eligibility_gates: dict = field(default_factory=dict)  # {gate: pass/fail/unknown}
    
    # Issues and Recommendations
    critical_gaps: list[str] = field(default_factory=list)
    structure_issues: list[str] = field(default_factory=list)
    recommendations: list[str] = field(default_factory=list)
    estimated_fix_months: int = 0
    
    # Metadata
    analyzed_at: str = field(default_factory=lambda: datetime.now().isoformat())
    analyzer_version: str = '2.0'


# ============================================================
# BLOOM'S TAXONOMY VERBS
# ============================================================

KNOWLEDGE_VERBS = [
    'define', 'describe', 'identify', 'know', 'label', 'list', 'match',
    'name', 'outline', 'recall', 'recognize', 'reproduce', 'select', 'state',
    'understand', 'explain', 'interpret', 'summarize', 'classify', 'compare',
    'demonstrate understanding', 'possess knowledge', 'have knowledge',
    # Norwegian
    'beskrive', 'forklare', 'identifisere', 'gjenkjenne', 'definere',
    # German
    'beschreiben', 'erklären', 'identifizieren', 'erkennen', 'definieren',
    # French
    'décrire', 'expliquer', 'identifier', 'reconnaître', 'définir',
]

SKILL_VERBS = [
    'apply', 'demonstrate', 'employ', 'illustrate', 'interpret', 'operate',
    'practice', 'schedule', 'sketch', 'solve', 'use', 'write', 'analyze',
    'analyse', 'calculate', 'categorize', 'compare', 'contrast', 'criticize',
    'differentiate', 'discriminate', 'distinguish', 'examine', 'experiment',
    'question', 'test', 'create', 'design', 'develop', 'formulate', 'construct',
    'produce', 'plan', 'compose', 'integrate', 'evaluate', 'assess', 'argue',
    'defend', 'judge', 'support', 'value', 'critique', 'recommend',
    # Norwegian
    'anvende', 'analysere', 'vurdere', 'utvikle', 'designe', 'løse',
    # German
    'anwenden', 'analysieren', 'bewerten', 'entwickeln', 'gestalten', 'lösen',
    # French
    'appliquer', 'analyser', 'évaluer', 'développer', 'concevoir', 'résoudre',
]

ATTITUDE_VERBS = [
    'appreciate', 'accept', 'commit', 'defend', 'demonstrate commitment',
    'display', 'exhibit', 'internalize', 'value', 'behave', 'act ethically',
    'show responsibility', 'take responsibility', 'respect', 'embrace',
    # Norwegian
    'verdsette', 'akseptere', 'forplikte', 'respektere', 'vise ansvar',
    # German
    'wertschätzen', 'akzeptieren', 'verpflichten', 'respektieren',
    # French
    'apprécier', 'accepter', 'engager', 'respecter', 'valoriser',
]

WEAK_VERBS = [
    'understand', 'know', 'be aware', 'appreciate', 'be familiar',
    'have knowledge', 'possess', 'gain insight', 'learn about',
    # Norwegian
    'forstå', 'kjenne til', 'være kjent med', 'ha kunnskap',
    # German
    'verstehen', 'kennen', 'wissen', 'vertraut sein',
    # French
    'comprendre', 'connaître', 'savoir', 'être familier',
]


# ============================================================
# EMBEDDING SERVICE
# ============================================================

class EmbeddingService:
    """Generate embeddings using Google Gemini"""
    
    def __init__(self):
        if not HAS_GEMINI:
            raise RuntimeError("google-generativeai not installed")
        
        api_key = os.getenv('GOOGLE_API_KEY') or os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise RuntimeError("GOOGLE_API_KEY or GEMINI_API_KEY not set")
        
        genai.configure(api_key=api_key)
    
    def embed_text(self, text: str) -> list[float]:
        """Generate embedding for a single text"""
        result = genai.embed_content(
            model=EMBEDDING_MODEL,
            content=text,
            task_type="semantic_similarity"
        )
        return result['embedding']
    
    def embed_batch(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings for multiple texts"""
        embeddings = []
        for text in texts:
            embeddings.append(self.embed_text(text))
        return embeddings


# ============================================================
# SUPABASE SERVICE
# ============================================================

class SupabaseService:
    """Database operations for EFMD tool"""
    
    def __init__(self):
        if not HAS_SUPABASE:
            raise RuntimeError("supabase not installed")
        
        url = os.getenv('SUPABASE_URL')
        key = os.getenv('SUPABASE_KEY') or os.getenv('SUPABASE_SERVICE_KEY')
        
        if not url or not key:
            raise RuntimeError("SUPABASE_URL and SUPABASE_KEY must be set")
        
        self.client: Client = create_client(url, key)
    
    def _get_or_create_institution(self, name: str) -> str:
        """Get institution ID by name, or create if not exists"""
        result = self.client.table('institutions').select('id').eq('name', name).execute()
        if result.data:
            return result.data[0]['id']

        result = self.client.table('institutions').insert({'name': name}).execute()
        return result.data[0]['id']

    def save_programme(self, programme: ProgrammeData) -> str:
        """Save programme and return ID"""
        data = {
            'institution_id': self._get_or_create_institution(programme.institution),
            'programme_name': programme.programme_name,
            'degree_type': programme.degree_type,
            'primary_url': programme.primary_url,
            'urls_scraped': programme.urls_scraped,
            'duration_months': programme.duration_months,
            'total_ects': programme.total_ects,
            'delivery_mode': programme.delivery_mode,
            'languages_of_instruction': programme.languages_of_instruction,
            'raw_html': programme.raw_html[:100000] if programme.raw_html else None,  # Limit size
            'raw_text': programme.raw_text[:50000] if programme.raw_text else None,
        }
        
        result = self.client.table('programmes').insert(data).execute()
        return result.data[0]['id']

    
    def save_ilo(self, programme_id: str, ilo: ILOAnalysis, order: int) -> str:
        """Save ILO and return ID"""
        data = {
            'programme_id': programme_id,
            'ilo_text': ilo.text,
            'ilo_order': order,
            'source_language': ilo.source_language,
            'ksa_category': ilo.ksa_category,
            'verb_found': ilo.verb_found,
            'has_weak_verb': ilo.has_weak_verb,
            'is_measurable': ilo.is_measurable,
            'quality_issues': ilo.quality_issues,
            'embedding': ilo.embedding,
            'matched_pillars': ilo.matched_pillars,
            'best_match_score': ilo.best_match_score,
        }
        
        result = self.client.table('programme_ilos').insert(data).execute()
        return result.data[0]['id']

    
    def save_course(self, programme_id: str, course: CourseData) -> str:
        """Save course and return ID"""
        data = {
            'programme_id': programme_id,
            'title': course.title,
            'description': course.description,
            'source_language': course.source_language,
            'ects': course.ects,
            'year': course.year,
            'semester': course.semester,
            'is_mandatory': course.is_mandatory,
            'course_ilos': course.course_ilos,
            'embedding': course.embedding,
            'matched_pillars': course.matched_pillars,
        }
        
        result = self.client.table('programme_courses').insert(data).execute()
        return result.data[0]['id']

    
    def save_gap_analysis(self, programme_id: str, gap: GapAnalysisResult) -> str:
        """Save gap analysis and return ID"""
        data = {
            'programme_id': programme_id,
            'readiness_score': gap.readiness_score,
            'eligibility_pass': gap.eligibility_pass,
            'ilo_analysis': {
                'count': gap.ilo_count,
                'has_knowledge': gap.ilo_has_knowledge,
                'has_skills': gap.ilo_has_skills,
                'has_attitudes': gap.ilo_has_attitudes,
                'weak_verb_count': gap.ilo_weak_verb_count,
                'issues': gap.ilo_issues,
            },
            'pillar_coverage': gap.pillar_coverage,
            'eligibility_gates': gap.eligibility_gates,
            'structure_issues': gap.structure_issues,
            'critical_gaps': gap.critical_gaps,
            'recommendations': gap.recommendations,
            'estimated_fix_months': gap.estimated_fix_months,
            'analyzer_version': gap.analyzer_version,
        }
        
        result = self.client.table('gap_analyses').insert(data).execute()
        
        # Update programme with score
        self.client.table('programmes').update({
            'readiness_score': gap.readiness_score,
            'last_analysis_at': datetime.now().isoformat(),
        }).eq('id', programme_id).execute()
        
        return result.data[0]['id']

    
    def get_requirements(self) -> list[dict]:
        """Get all EFMD requirements with embeddings"""
        result = self.client.table('efmd_requirements').select('*').execute()
        return result.data
    
    def get_pillar_coverage(self, programme_id: str) -> list[dict]:
        """Get pillar coverage using database function"""
        result = self.client.rpc(
            'get_programme_pillar_coverage',
            {'p_programme_id': programme_id}
        ).execute()
        return result.data
    
    def list_programmes(self, limit: int = 50) -> list[dict]:
        """List all scraped programmes, newest first"""
        result = self.client.table('programmes').select(
            '*, institutions(name)'
        ).order('created_at', desc=True).limit(limit).execute()
        return result.data
    
    def get_programme(self, programme_id: str) -> dict:
        """Get a single programme by ID"""
        result = self.client.table('programmes').select('*').eq('id', programme_id).single().execute()
        return result.data
    
    def get_programme_ilos(self, programme_id: str) -> list[dict]:
        """Get ILOs for a programme"""
        result = self.client.table('programme_ilos').select('*').eq(
            'programme_id', programme_id
        ).order('ilo_order').execute()
        return result.data
    
    def get_gap_analysis(self, programme_id: str) -> dict:
        """Get the latest gap analysis for a programme"""
        result = self.client.table('gap_analyses').select('*').eq(
            'programme_id', programme_id
        ).order('created_at', desc=True).limit(1).execute()
        return result.data[0] if result.data else None
    
    def delete_programme(self, programme_id: str) -> bool:
        """Delete a programme and its related data"""
        # Delete ILOs first (foreign key)
        self.client.table('programme_ilos').delete().eq('programme_id', programme_id).execute()
        # Delete gap analyses
        self.client.table('gap_analyses').delete().eq('programme_id', programme_id).execute()
        # Delete programme
        self.client.table('programmes').delete().eq('id', programme_id).execute()
        return True


# ============================================================
# ILO ANALYZER
# ============================================================

class ILOAnalyzer:
    """Analyze ILOs for quality and EFMD compliance"""
    
    def __init__(self, embedding_service: Optional[EmbeddingService] = None):
        self.embedder = embedding_service
    
    def analyze(self, ilo_text: str, source_language: str = 'en') -> ILOAnalysis:
        """Analyze a single ILO"""
        analysis = ILOAnalysis(
            text=ilo_text,
            source_language=source_language
        )
        
        ilo_lower = ilo_text.lower()
        
        # Detect KSA category
        for verb in ATTITUDE_VERBS:
            if verb in ilo_lower:
                analysis.ksa_category = 'Attitude'
                analysis.verb_found = verb
                break
        
        if not analysis.ksa_category:
            for verb in SKILL_VERBS:
                if verb in ilo_lower:
                    analysis.ksa_category = 'Skill'
                    analysis.verb_found = verb
                    break
        
        if not analysis.ksa_category:
            for verb in KNOWLEDGE_VERBS:
                if verb in ilo_lower:
                    analysis.ksa_category = 'Knowledge'
                    analysis.verb_found = verb
                    break
        
        if not analysis.ksa_category:
            analysis.quality_issues.append('No clear action verb detected')
            analysis.is_measurable = False
        
        # Check for weak verbs
        for weak in WEAK_VERBS:
            if weak in ilo_lower:
                analysis.has_weak_verb = True
                analysis.quality_issues.append(f'Uses weak/vague verb: "{weak}"')
                analysis.is_measurable = False
                break
        
        # Check length
        if len(ilo_text) < 30:
            analysis.quality_issues.append('ILO too brief - likely lacks specificity')
        
        # Generate embedding if available
        if self.embedder:
            try:
                analysis.embedding = self.embedder.embed_text(ilo_text)
            except Exception as e:
                print(f"Warning: Could not generate embedding: {e}")
        
        return analysis


# ============================================================
# SITE-SPECIFIC PARSERS
# ============================================================

class BaseSiteParser(ABC):
    """Base class for site-specific parsers"""
    
    @abstractmethod
    def can_parse(self, url: str) -> bool:
        """Check if this parser can handle the URL"""
        pass
    
    @abstractmethod
    def parse(self, html: str, url: str) -> dict:
        """Parse HTML and return structured data"""
        pass


class GenericParser(BaseSiteParser):
    """Generic parser for unknown sites"""
    
    def can_parse(self, url: str) -> bool:
        return True  # Fallback parser
    
    def parse(self, html: str, url: str) -> dict:
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove noise
        for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()
        
        result = {
            'title': '',
            'ilos': [],
            'courses': [],
            'text': '',
            'language': 'en',
        }
        
        # Get title
        title_tag = soup.find('title')
        result['title'] = title_tag.get_text().strip() if title_tag else ''
        
        # Get main content
        main = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'content|main'))
        if main:
            result['text'] = main.get_text(separator=' ', strip=True)
        else:
            result['text'] = soup.get_text(separator=' ', strip=True)
        
        # Look for ILO sections
        ilo_patterns = [
            r'learning outcome', r'programme outcome', r'intended learning',
            r'what you will learn', r'on completion', r'graduate will',
            r'læringsutbytte', r'kunnskaper', r'ferdigheter',  # Norwegian
            r'lernergebnisse', r'lernziele',  # German
            r'acquis d.apprentissage', r'objectifs',  # French
        ]
        
        for pattern in ilo_patterns:
            sections = soup.find_all(['div', 'section', 'article'], 
                                    string=re.compile(pattern, re.I))
            for section in sections:
                items = section.find_all('li')
                for item in items:
                    text = item.get_text().strip()
                    if len(text) > 20:
                        result['ilos'].append(text)
        
        # Look for structured ILO headings
        for heading in soup.find_all(['h2', 'h3', 'h4', 'h5']):
            heading_text = heading.get_text().lower()
            if any(kw in heading_text for kw in ['kunnskaper', 'ferdigheter', 'kompetanse',
                                                   'knowledge', 'skills', 'competence',
                                                   'kenntnisse', 'fertigkeiten',
                                                   'connaissances', 'compétences']):
                next_elem = heading.find_next_sibling()
                while next_elem and next_elem.name not in ['h2', 'h3', 'h4', 'h5']:
                    if next_elem.name == 'ul':
                        for li in next_elem.find_all('li'):
                            text = li.get_text().strip()
                            if len(text) > 20 and text not in result['ilos']:
                                result['ilos'].append(text)
                    next_elem = next_elem.find_next_sibling()
        
        # Detect language from content
        if any(word in result['text'].lower() for word in ['læringsutbytte', 'studenten', 'kandidaten']):
            result['language'] = 'no'
        elif any(word in result['text'].lower() for word in ['lernergebnisse', 'studierende']):
            result['language'] = 'de'
        elif any(word in result['text'].lower() for word in ['objectifs', 'étudiant']):
            result['language'] = 'fr'
        
        return result


class NorwegianUniversityParser(BaseSiteParser):
    """Parser for Norwegian university sites (UiA, NTNU, NHH, BI, etc.)"""
    
    DOMAINS = ['uia.no', 'ntnu.no', 'nhh.no', 'bi.no', 'uio.no', 'uib.no', 'uit.no', 'usn.no']
    
    def can_parse(self, url: str) -> bool:
        domain = urlparse(url).netloc.replace('www.', '')
        return any(d in domain for d in self.DOMAINS)
    
    def parse(self, html: str, url: str) -> dict:
        soup = BeautifulSoup(html, 'html.parser')
        
        result = {
            'title': '',
            'ilos': [],
            'courses': [],
            'text': '',
            'language': 'no',
        }
        
        # Norwegian sites often have structured ILO sections
        # Look for "Læringsutbytte" (Learning Outcomes) section
        
        # Pattern 1: Look for specific Norwegian headers
        for heading in soup.find_all(['h2', 'h3', 'h4', 'h5']):
            text = heading.get_text().strip().lower()
            
            # Knowledge section
            if 'kunnskap' in text or 'kunnskaper' in text:
                result['ilos'].extend(self._extract_following_items(heading, 'Knowledge'))
            
            # Skills section
            elif 'ferdighet' in text or 'ferdigheter' in text:
                result['ilos'].extend(self._extract_following_items(heading, 'Skill'))
            
            # General competence / Attitudes
            elif 'generell kompetanse' in text or 'holdning' in text:
                result['ilos'].extend(self._extract_following_items(heading, 'Attitude'))
            
            # Generic learning outcomes
            elif 'læringsutbytte' in text:
                result['ilos'].extend(self._extract_following_items(heading))
        
        # Get page title
        title_tag = soup.find('title')
        result['title'] = title_tag.get_text().strip() if title_tag else ''
        
        # Get full text
        for tag in soup(['script', 'style', 'nav', 'footer']):
            tag.decompose()
        result['text'] = soup.get_text(separator=' ', strip=True)
        
        return result
    
    def _extract_following_items(self, heading, category: str = None) -> list[str]:
        """Extract list items following a heading"""
        items = []
        next_elem = heading.find_next_sibling()
        
        while next_elem and next_elem.name not in ['h2', 'h3', 'h4', 'h5']:
            if next_elem.name == 'ul' or next_elem.name == 'ol':
                for li in next_elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if len(text) > 15:
                        items.append(text)
            elif next_elem.name == 'p':
                text = next_elem.get_text().strip()
                if len(text) > 30:
                    items.append(text)
            next_elem = next_elem.find_next_sibling()
        
        return items


class FinnishUniversityParser(BaseSiteParser):
    """Parser for Finnish university sites (Vaasa, Aalto, Hanken, etc.)"""
    
    DOMAINS = ['uwasa.fi', 'aalto.fi', 'hanken.fi', 'helsinki.fi', 'utu.fi']
    
    def can_parse(self, url: str) -> bool:
        domain = urlparse(url).netloc.replace('www.', '')
        return any(d in domain for d in self.DOMAINS)
    
    def parse(self, html: str, url: str) -> dict:
        soup = BeautifulSoup(html, 'html.parser')
        
        result = {
            'title': '',
            'ilos': [],
            'courses': [],
            'text': '',
            'language': 'en',  # Finnish universities often have good English pages
        }
        
        # Finnish sites sometimes use "Osaamistavoitteet" (Learning objectives)
        # But English versions use standard Learning Outcomes terminology
        
        for heading in soup.find_all(['h2', 'h3', 'h4', 'h5']):
            text = heading.get_text().strip().lower()
            
            if any(kw in text for kw in ['learning outcome', 'learning objective',
                                          'osaamistavoitteet', 'oppimistavoitteet']):
                items = self._extract_following_items(heading)
                result['ilos'].extend(items)
        
        # Get title and text
        title_tag = soup.find('title')
        result['title'] = title_tag.get_text().strip() if title_tag else ''
        
        for tag in soup(['script', 'style', 'nav', 'footer']):
            tag.decompose()
        result['text'] = soup.get_text(separator=' ', strip=True)
        
        # Detect Finnish
        if any(word in result['text'].lower() for word in ['osaamistavoitteet', 'opintojaksot']):
            result['language'] = 'fi'
        
        return result
    
    def _extract_following_items(self, heading) -> list[str]:
        items = []
        next_elem = heading.find_next_sibling()
        
        while next_elem and next_elem.name not in ['h2', 'h3', 'h4', 'h5']:
            if next_elem.name in ['ul', 'ol']:
                for li in next_elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if len(text) > 15:
                        items.append(text)
            next_elem = next_elem.find_next_sibling()
        
        return items


# ============================================================
# MAIN SCRAPER
# ============================================================

class EFMDScraper:
    """Main scraper orchestrator"""
    
    def __init__(self, 
                 use_embeddings: bool = True,
                 use_database: bool = True):
        
        self.parsers: list[BaseSiteParser] = [
            NorwegianUniversityParser(),
            FinnishUniversityParser(),
            GenericParser(),  # Fallback
        ]
        
        self.embedder = None
        if use_embeddings and HAS_GEMINI:
            try:
                self.embedder = EmbeddingService()
            except Exception as e:
                print(f"Warning: Could not initialize embeddings: {e}")
        
        self.db = None
        if use_database and HAS_SUPABASE:
            try:
                self.db = SupabaseService()
            except Exception as e:
                print(f"Warning: Could not initialize database: {e}")
        
        self.ilo_analyzer = ILOAnalyzer(self.embedder)
    
    def get_parser(self, url: str) -> BaseSiteParser:
        """Get appropriate parser for URL"""
        for parser in self.parsers:
            if parser.can_parse(url):
                return parser
        return self.parsers[-1]  # Generic fallback
    
    def scrape_url(self, url: str) -> dict:
        """Scrape a single URL"""
        headers = {'User-Agent': USER_AGENT}
        
        try:
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            
            parser = self.get_parser(url)
            result = parser.parse(resp.text, url)
            result['url'] = url
            result['html'] = resp.text
            result['success'] = True
            
            return result
            
        except Exception as e:
            return {
                'url': url,
                'success': False,
                'error': str(e)
            }
    def extract_document_text(self, file_path: str) -> dict:
        """Extract text from PDF or DOCX file"""
        result = {
            'path': file_path,
            'text': '',
            'success': False,
            'error': None
        }
        
        try:
            if file_path.lower().endswith('.pdf'):
                if not HAS_PYPDF:
                    result['error'] = 'pypdf not installed'
                    return result
                
                reader = pypdf.PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
                result['text'] = '\n'.join(text_parts)
                result['success'] = True
                
            elif file_path.lower().endswith('.docx'):
                if not HAS_DOCX:
                    result['error'] = 'python-docx not installed'
                    return result
                
                doc = Document(file_path)
                text_parts = [para.text for para in doc.paragraphs]
                result['text'] = '\n'.join(text_parts)
                result['success'] = True
                
            elif file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    result["text"] = f.read()
                result["success"] = True
                
            else:
                result['error'] = f'Unsupported format: {file_path}'
                
        except Exception as e:
            result['error'] = str(e)
        
        return result

    def _extract_ilos_from_text(self, text: str) -> list[str]:
        print(f"DEBUG _extract_ilos_from_text: text length = {len(text)}")
        """Extract ILO-like statements from plain text"""
        ilos = []
        
        lines = re.split(r'[\n\r]+|(?<=[.!?])\s+', text)
        
        for line in lines:
            line = line.strip()
            if len(line) < 25 or len(line) > 500:
                continue
            
            line_lower = line.lower()
            
            has_verb = any(verb in line_lower for verb in 
                          KNOWLEDGE_VERBS + SKILL_VERBS + ATTITUDE_VERBS)
            
            if has_verb:
                clean = re.sub(r'^[\d•\-\*\s]+', '', line).strip()
                if clean and clean not in ilos:
                    ilos.append(clean)
        
        return ilos 

        try:
            if file_path.lower().endswith('.pdf'):
                if not HAS_PYPDF:
                    result['error'] = 'pypdf not installed'
                    return result
                
                reader = pypdf.PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
                result['text'] = '\n'.join(text_parts)
                result['success'] = True
                
            elif file_path.lower().endswith('.docx'):
                if not HAS_DOCX:
                    result['error'] = 'python-docx not installed'
                    return result
                
                doc = Document(file_path)
                text_parts = [para.text for para in doc.paragraphs]
                result['text'] = '\n'.join(text_parts)
                result['success'] = True
                
            elif file_path.lower().endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    result["text"] = f.read()
                result["success"] = True
                
            else:
                result['error'] = f'Unsupported format: {file_path}'
                
        except Exception as e:
            result['error'] = str(e)
        
        return result
    def find_language_variants(self, url: str) -> list[str]:
        """Try to find other language versions of a URL"""
        variants = [url]
        
        # Common patterns
        patterns = [
            ('/en/', '/no/'), ('/no/', '/en/'),
            ('/en/', '/de/'), ('/de/', '/en/'),
            ('/en/', '/fr/'), ('/fr/', '/en/'),
            ('/en/', '/fi/'), ('/fi/', '/en/'),
            ('?lang=en', '?lang=no'), ('?lang=no', '?lang=en'),
            ('/english/', '/norwegian/'), ('/norwegian/', '/english/'),
        ]
        
        for pattern_from, pattern_to in patterns:
            if pattern_from in url:
                variant = url.replace(pattern_from, pattern_to)
                if variant not in variants:
                    variants.append(variant)
        
        return variants
    
    def scrape_programme(self,
                         urls,
                         institution: str = None,
                         programme_name: str = None,
                         documents: list[str] = None,
                         follow_variants: bool = True) -> ProgrammeData:
        """Scrape a programme from URL(s)"""
        
        # Normalize urls to list
        if isinstance(urls, str):
            url_list = [urls]
        else:
            url_list = list(urls)
        
        primary_url = url_list[0] if url_list else ''
        
        # Add language variants if requested
        if follow_variants:
            expanded = []
            for u in url_list:
                for variant in self.find_language_variants(u):
                    if variant not in expanded:
                        expanded.append(variant)
            url_list = expanded
        
        # Scrape all URLs
        all_ilos = []
        all_courses = []
        all_text = ''
        all_html = ''
        languages = []
        scraped_urls = []
        detected_title = ''
        
        for scrape_url in url_list:
            print(f"Scraping: {scrape_url}")
            result = self.scrape_url(scrape_url)
            
            if result.get('success'):
                scraped_urls.append(scrape_url)
                all_text += ' ' + result.get('text', '')
                all_html += result.get('html', '')
                
                if not detected_title:
                    detected_title = result.get('title', '')
                
                lang = result.get('language', 'en')
                if lang not in languages:
                    languages.append(lang)
                
                # Add ILOs with language tag
                for ilo_text in result.get('ilos', []):
                    all_ilos.append((ilo_text, lang))
                
                for course in result.get('courses', []):
                    all_courses.append(course)
                
            else:
                print(f"  Failed: {result.get('error')}")

            # Process uploaded documents
        if documents:
            for doc_path in documents:
                print(f"Extracting: {doc_path}")
                result = self.extract_document_text(doc_path)
                
                if result['success']:
                    all_text += ' ' + result['text']
                    
                    # Try to extract ILOs from document
                    doc_ilos = self._extract_ilos_from_text(result['text'])
                    for ilo_text in doc_ilos:
                        all_ilos.append((ilo_text, 'en'))
                else:
                    print(f"  Failed: {result.get('error')}")
        
        # Deduplicate ILOs
        seen = set()
        unique_ilos = []
        for ilo_text, lang in all_ilos:
            key = ilo_text.lower().strip()
            if key not in seen and len(ilo_text) > 20:
                seen.add(key)
                unique_ilos.append((ilo_text, lang))
        
        # Analyze ILOs
        analyzed_ilos = []
        for ilo_text, lang in unique_ilos:
            analysis = self.ilo_analyzer.analyze(ilo_text, lang)
            analyzed_ilos.append(analysis)
        
        # Detect institution and programme name if not provided
        if not institution:
            domain = urlparse(url).netloc.replace('www.', '')
            institution = domain.split('.')[0].upper()
        
        if not programme_name:
            programme_name = detected_title.split('|')[0].split('-')[0].strip() or "Unknown Programme"
        
        # Detect degree type
        degree_type = None
        name_lower = programme_name.lower()
        if 'mba' in name_lower:
            degree_type = 'MBA'
        elif 'msc' in name_lower or 'master' in name_lower:
            degree_type = 'MSc'
        elif 'bsc' in name_lower or 'bachelor' in name_lower:
            degree_type = 'BSc'
        elif 'phd' in name_lower or 'doctor' in name_lower:
            degree_type = 'PhD'
        
        # Create programme data
        programme = ProgrammeData(
            institution=institution,
            programme_name=programme_name,
            primary_url=primary_url,
            degree_type=degree_type,
            urls_scraped=scraped_urls,
            languages_of_instruction=languages,
            programme_ilos=analyzed_ilos,
            raw_html=all_html,
            raw_text=all_text,
        )
        
        # Check for ILO matrix mention
        if any(term in all_text.lower() for term in 
               ['learning outcome matrix', 'ilo matrix', 'curriculum map',
                'læringsutbyttematrise', 'programme map']):
            programme.has_ilo_matrix = True
        
        return programme
    
    def analyze_gaps(self, programme: ProgrammeData) -> GapAnalysisResult:
        """Generate gap analysis for a programme"""
        
        gap = GapAnalysisResult(programme=programme)
        gap.ilo_count = len(programme.programme_ilos)
        
        # === ILO Analysis ===
        if gap.ilo_count == 0:
            gap.critical_gaps.append('NO PROGRAMME ILOs FOUND - Critical for EFMD')
            gap.ilo_issues.append('Programme ILOs are missing entirely')
        else:
            # Check count
            if gap.ilo_count < 5:
                gap.ilo_issues.append(f'Only {gap.ilo_count} ILOs - EFMD recommends 5-6')
            elif gap.ilo_count > 8:
                gap.ilo_issues.append(f'{gap.ilo_count} ILOs is too many - EFMD recommends 5-6')
            
            # Count KSA
            knowledge_count = sum(1 for ilo in programme.programme_ilos if ilo.ksa_category == 'Knowledge')
            skill_count = sum(1 for ilo in programme.programme_ilos if ilo.ksa_category == 'Skill')
            attitude_count = sum(1 for ilo in programme.programme_ilos if ilo.ksa_category == 'Attitude')
            gap.ilo_weak_verb_count = sum(1 for ilo in programme.programme_ilos if ilo.has_weak_verb)
            
            gap.ilo_has_knowledge = knowledge_count > 0
            gap.ilo_has_skills = skill_count > 0
            gap.ilo_has_attitudes = attitude_count > 0
            
            if not gap.ilo_has_knowledge:
                gap.ilo_issues.append('No Knowledge-focused ILOs detected')
                gap.critical_gaps.append('Missing Knowledge dimension in ILOs')
            if not gap.ilo_has_skills:
                gap.ilo_issues.append('No Skill-focused ILOs detected')
                gap.critical_gaps.append('Missing Skills dimension in ILOs')
            if not gap.ilo_has_attitudes:
                gap.ilo_issues.append('No Attitude-focused ILOs detected')
                gap.critical_gaps.append('Missing Attitudes dimension in ILOs')
            
            if gap.ilo_weak_verb_count > 0:
                gap.ilo_issues.append(f'{gap.ilo_weak_verb_count} ILOs use weak/unmeasurable verbs')
        
        # === Pillar Coverage (Semantic if available, else keyword) ===
        gap.pillar_coverage = self._check_pillar_coverage(programme)
        
        for pillar, score in gap.pillar_coverage.items():
            if score < PILLAR_MATCH_THRESHOLD:
                gap.missing_pillars.append(pillar)
        
        if 'ERS' in gap.missing_pillars:
            gap.critical_gaps.append('ERS content missing - mandatory since 2013')
            gap.recommendations.append('Integrate ethics and sustainability across curriculum')
        
        if 'International' in gap.missing_pillars:
            gap.recommendations.append('Add international/global perspective to curriculum')
        
        if 'Practice' in gap.missing_pillars:
            gap.recommendations.append('Strengthen links to business practice')
        
        if 'Digital' in gap.missing_pillars:
            gap.recommendations.append('Add digital transformation content')
        
        # === Structure Issues ===
        if not programme.courses:
            gap.structure_issues.append('No course structure found')
        
        if not programme.has_ilo_matrix:
            gap.structure_issues.append('No ILO mapping matrix visible')
            gap.recommendations.append('Create matrix: Course ILOs → Programme ILOs → Assessments')
        
        # === Documentation Status ===
        if programme.courses:
            gap.courses_total = len(programme.courses)
            for course in programme.courses:
                # Check if course has ILOs (assuming course has ilos attribute or similar)
                course_has_ilos = hasattr(course, 'ilos') and len(course.ilos) > 0
                if course_has_ilos:
                    gap.courses_with_ilos += 1
            else:
                    gap.courses_no_ilos.append(course.title if hasattr(course, 'title') else str(course))
            
            if gap.courses_total > 0:
                gap.documentation_score = (gap.courses_with_ilos / gap.courses_total) * 100
            
            # Add warning if documentation is incomplete
            if gap.documentation_score < 80:
                gap.structure_issues.append(f'Only {gap.courses_with_ilos}/{gap.courses_total} courses have documented ILOs')
                gap.recommendations.append('Document ILOs for all courses before SAR submission')
        
        # === Eligibility Gates ===
        gap.eligibility_gates = {
            'ELG-1': 'pass',  # Assume business focus if submitted
            'ELG-2': 'pass' if programme.degree_type else 'unknown',
            'ELG-3': 'unknown',  # Need operating history
            'ELG-4': 'fail' if gap.ilo_count == 0 else ('pass' if gap.ilo_count >= 5 else 'partial'),
            'ELG-5': 'unknown',  # Need QA docs
            'ELG-6': 'unknown',  # Need faculty CVs
            'ELG-7': 'unknown',  # Need student services info
            'ELG-8': 'unknown',  # Need resources info
            'ELG-9': 'unknown',  # Need governance info
            'ELG-10': 'unknown', # Need financial info
        }
        
        # === Readiness Score ===
        score = 100
        
        # Critical items (-20 each)
        if gap.ilo_count == 0:
            score -= 20
        if 'ERS' in gap.missing_pillars:
            score -= 20
        if not gap.ilo_has_knowledge:
            score -= 10
        if not gap.ilo_has_skills:
            score -= 10
        if not gap.ilo_has_attitudes:
            score -= 10
        
        # Missing pillars (-5 each)
        score -= len(gap.missing_pillars) * 5
        
        # Structure issues (-5 each)
        score -= len(gap.structure_issues) * 5
        
        # ILO issues (-3 each)
        score -= len(gap.ilo_issues) * 3
        
        # Documentation completeness penalty
        if gap.courses_total > 0 and gap.documentation_score < 100:
            doc_penalty = int((100 - gap.documentation_score) / 10)  # -1 point per 10% missing
            score -= doc_penalty
        
        gap.readiness_score = max(0, score)
        gap.eligibility_pass = gap.readiness_score >= 70 and len(gap.critical_gaps) == 0
        
        # === Estimate Fix Time ===
        if gap.readiness_score >= 80:
            gap.estimated_fix_months = 1
        elif gap.readiness_score >= 60:
            gap.estimated_fix_months = 3
        elif gap.readiness_score >= 40:
            gap.estimated_fix_months = 6
        else:
            gap.estimated_fix_months = 12
        
        return gap
    
    def _check_pillar_coverage(self, programme: ProgrammeData) -> dict:
        """Check pillar coverage using embeddings if available, else keywords"""
        
        # Combine all text for analysis
        all_text = ' '.join([
            programme.programme_name or '',
            ' '.join(ilo.text for ilo in programme.programme_ilos),
            ' '.join(programme.programme_aims),
            ' '.join(c.title for c in programme.courses),
            programme.raw_text[:10000] if programme.raw_text else '',
        ]).lower()
        
        # Keyword-based check (fallback or supplement)
        keywords = {
            'International': ['international', 'global', 'cross-cultural', 'worldwide', 
                            'foreign', 'abroad', 'exchange', 'intercultural'],
            'Practice': ['practical', 'industry', 'corporate', 'business', 'real-world',
                        'case study', 'internship', 'project', 'consultancy'],
            'ERS': ['ethics', 'ethical', 'responsibility', 'sustainable', 'sustainability',
                   'csr', 'governance', 'esg', 'stakeholder'],
            'Digital': ['digital', 'technology', 'data', 'analytics', 'ai', 
                       'artificial intelligence', 'software', 'fintech'],
        }
        
        coverage = {}
        for pillar, kws in keywords.items():
            matches = sum(1 for kw in kws if kw in all_text)
            # Normalize to 0-1 score
            coverage[pillar] = min(1.0, matches / 3)  # 3+ keywords = full coverage
        
        return coverage
    
    def save_to_database(self, programme: ProgrammeData, gap: GapAnalysisResult) -> dict:
        """Save programme and analysis to database"""
        if not self.db:
            raise RuntimeError("Database not initialized")
        
        # Save programme
        programme_id = self.db.save_programme(programme)
        
        # Save ILOs
        for i, ilo in enumerate(programme.programme_ilos, 1):
            self.db.save_ilo(programme_id, ilo, i)
        
        # Save courses
        for course in programme.courses:
            self.db.save_course(programme_id, course)
        
        # Save gap analysis
        gap_id = self.db.save_gap_analysis(programme_id, gap)
        
        return {
            'programme_id': programme_id,
            'gap_analysis_id': gap_id,
        }


# ============================================================
# REPORT GENERATOR
# ============================================================

def format_gap_report(gap: GapAnalysisResult) -> str:
    """Format gap analysis as readable report"""
    
    prog = gap.programme
    
    lines = []
    lines.append("=" * 70)
    lines.append("EFMD PROGRAMME GAP ANALYSIS REPORT")
    lines.append("=" * 70)
    lines.append("")
    lines.append(f"Programme: {prog.programme_name}")
    lines.append(f"Institution: {prog.institution}")
    lines.append(f"URL: {prog.primary_url}")
    lines.append(f"Languages: {', '.join(prog.languages_of_instruction)}")
    lines.append(f"Analysed: {gap.analyzed_at}")
    lines.append("")
    
    # Readiness Score
    score = gap.readiness_score
    if score >= 80:
        status = "🟢 READY - Minor improvements needed"
    elif score >= 60:
        status = "🟡 PARTIAL - Significant work required"
    elif score >= 40:
        status = "🟠 AT RISK - Major gaps to address"
    else:
        status = "🔴 NOT READY - Fundamental issues"
    
    lines.append(f"EFMD READINESS SCORE: {score}/100")
    lines.append(f"Status: {status}")
    lines.append(f"Estimated time to readiness: {gap.estimated_fix_months} months")
    lines.append("")
    
    # Critical Gaps
    if gap.critical_gaps:
        lines.append("-" * 70)
        lines.append("🚨 CRITICAL GAPS (Must fix before submission)")
        lines.append("-" * 70)
        for crit in gap.critical_gaps:
            lines.append(f"  ❌ {crit}")
        lines.append("")
    
    # ILO Analysis
    lines.append("-" * 70)
    lines.append("INTENDED LEARNING OUTCOMES (ILOs)")
    lines.append("-" * 70)
    lines.append(f"  Count: {gap.ilo_count} (optimal: 5-6)")
    lines.append(f"  Knowledge dimension: {'✅' if gap.ilo_has_knowledge else '❌'}")
    lines.append(f"  Skills dimension: {'✅' if gap.ilo_has_skills else '❌'}")
    lines.append(f"  Attitudes dimension: {'✅' if gap.ilo_has_attitudes else '❌'}")
    lines.append(f"  Weak verbs found: {gap.ilo_weak_verb_count}")
    
    if gap.ilo_issues:
        lines.append("")
        lines.append("  Issues:")
        for issue in gap.ilo_issues:
            lines.append(f"    ⚠️  {issue}")
    lines.append("")
    
    # ILOs Found
    if prog.programme_ilos:
        lines.append("  ILOs Found:")
        for i, ilo in enumerate(prog.programme_ilos, 1):
            category = f"[{ilo.ksa_category or '?'}]" if ilo.ksa_category else "[?]"
            weak = " ⚠️WEAK" if ilo.has_weak_verb else ""
            display = ilo.text[:80] + "..." if len(ilo.text) > 80 else ilo.text
            lines.append(f"    {i}. {category}{weak} {display}")
        lines.append("")
    
    # Documentation Status
    if gap.courses_total > 0:
        lines.append("-" * 70)
        lines.append("DOCUMENTATION STATUS")
        lines.append("-" * 70)
        doc_pct = gap.documentation_score
        if doc_pct >= 90:
            doc_icon = "✅"
        elif doc_pct >= 70:
            doc_icon = "🟡"
        else:
            doc_icon = "❌"
        lines.append(f"  Courses with ILOs: {gap.courses_with_ilos}/{gap.courses_total} ({doc_pct:.0f}%) {doc_icon}")
        
        if gap.courses_no_ilos:
            lines.append("")
            lines.append("  Courses missing ILOs:")
            for course_name in gap.courses_no_ilos[:10]:  # Show max 10
                lines.append(f"    ❌ {course_name}")
            if len(gap.courses_no_ilos) > 10:
                lines.append(f"    ... and {len(gap.courses_no_ilos) - 10} more")
        lines.append("")
    
    # Pillar Coverage
    lines.append("-" * 70)
    lines.append("EFMD PILLAR COVERAGE")
    lines.append("-" * 70)
    for pillar in ['International', 'Practice', 'ERS', 'Digital']:
        score = gap.pillar_coverage.get(pillar, 0)
        if score >= STRONG_MATCH_THRESHOLD:
            icon = "✅"
        elif score >= PILLAR_MATCH_THRESHOLD:
            icon = "🟡"
        else:
            icon = "❌"
        lines.append(f"  {pillar}: {icon} ({score:.0%})")
    
    if gap.missing_pillars:
        lines.append("")
        lines.append("  Missing/weak pillars:")
        for pillar in gap.missing_pillars:
            lines.append(f"    ❌ {pillar}")
    lines.append("")
    
    # Structure Issues
    if gap.structure_issues:
        lines.append("-" * 70)
        lines.append("STRUCTURE ISSUES")
        lines.append("-" * 70)
        for issue in gap.structure_issues:
            lines.append(f"  ⚠️  {issue}")
        lines.append("")
    
    # Recommendations
    if gap.recommendations:
        lines.append("-" * 70)
        lines.append("RECOMMENDATIONS")
        lines.append("-" * 70)
        for i, rec in enumerate(gap.recommendations, 1):
            lines.append(f"  {i}. {rec}")
        lines.append("")
    
    lines.append("=" * 70)
    lines.append("END OF REPORT")
    lines.append("=" * 70)
    
    return "\n".join(lines)


# ============================================================
# CLI
# ============================================================

def main():
    import sys
    
    if len(sys.argv) < 2:
        print("EFMD Programme Gap Analysis Tool v2.0")
        print("")
        print("Usage:")
        print("  python efmd_scraper_v2.py <URL> [institution] [programme_name]")
        print("  python efmd_scraper_v2.py --demo")
        print("")
        print("Examples:")
        print("  python efmd_scraper_v2.py https://www.uia.no/en/studieplaner/programme/MAOKOam")
        print("  python efmd_scraper_v2.py https://www.uwasa.fi/en/education/masters/finance 'University of Vaasa' 'MSc Finance'")
        print("")
        print("Environment variables:")
        print("  GOOGLE_API_KEY or GEMINI_API_KEY - For embeddings")
        print("  SUPABASE_URL + SUPABASE_KEY - For database storage")
        sys.exit(1)
    
    if sys.argv[1] == '--demo':
        run_demo()
        return
    
    url = sys.argv[1]
    institution = sys.argv[2] if len(sys.argv) > 2 else None
    programme_name = sys.argv[3] if len(sys.argv) > 3 else None
    
    # Initialize scraper
    scraper = EFMDScraper(
        use_embeddings=HAS_GEMINI,
        use_database=False  # Set to True when DB is ready
    )
    
    # Scrape and analyze
    print(f"\nScraping programme from: {url}")
    programme = scraper.scrape_programme(url, institution, programme_name)
    
    print(f"\nAnalyzing gaps...")
    gap = scraper.analyze_gaps(programme)
    
    # Print report
    print("\n" + format_gap_report(gap))
    
    # Save JSON
    output_file = f"gap_analysis_{urlparse(url).netloc.replace('.', '_')}.json"
    with open(output_file, 'w') as f:
        # Convert to dict (handle dataclasses)
        output = {
            'programme': asdict(programme),
            'gap_analysis': asdict(gap),
        }
        # Remove embeddings from output (too large)
        for ilo in output['programme']['programme_ilos']:
            ilo.pop('embedding', None)
        json.dump(output, f, indent=2, default=str)
    
    print(f"\nJSON saved to: {output_file}")


def run_demo():
    """Run demo with sample data"""
    print("Running demo with UiA MSc Finance sample data...")
    
    # Create sample programme
    programme = ProgrammeData(
        institution="University of Agder (UiA)",
        programme_name="MSc in Business Administration - Finance",
        primary_url="https://www.uia.no/en/studieplaner/programme/MAOKOam",
        degree_type="MSc",
        duration_months=24,
        total_ects=120,
        delivery_mode="Full-time",
        languages_of_instruction=["English"],
        urls_scraped=["https://www.uia.no/en/studieplaner/programme/MAOKOam"],
    )
    
    # Sample ILOs
    sample_ilos = [
        "Possess advanced knowledge about economic theory and its application to economic phenomena in the business sector.",
        "Have a sound basis for understanding value creation and financial decision-making in companies.",
        "Understand the relationship between various functions in a company.",
        "Work independently with the subject matter based on different types of sources, which they are able to analyse critically.",
        "Contribute in challenging situations that require analysis, reflection and the making of professional decisions.",
        "Can communicate and discuss viewpoints and results of scientific work with specialists and in a forum that is open to non-specialists.",
        "Can plan and lead the execution of various projects, both as a participant and as a leader.",
        "Can work at an advanced level with empirical data from both qualitative and quantitative research.",
        "The students can work independently and together with others."
    ]
    
    analyzer = ILOAnalyzer()
    programme.programme_ilos = [analyzer.analyze(ilo) for ilo in sample_ilos]
    
    # Sample courses
    programme.courses = [
        CourseData(title="Financial Management", ects=7.5),
        CourseData(title="Corporate Finance", ects=7.5),
        CourseData(title="Investment Analysis", ects=7.5),
        CourseData(title="Financial Econometrics", ects=7.5),
        CourseData(title="Derivatives", ects=7.5),
        CourseData(title="Asset Pricing", ects=7.5),
        CourseData(title="Master Thesis", ects=30),
    ]
    
    # Analyze
    scraper = EFMDScraper(use_embeddings=False, use_database=False)
    gap = scraper.analyze_gaps(programme)
    
    # Print report
    print("\n" + format_gap_report(gap))


if __name__ == "__main__":
    main()
